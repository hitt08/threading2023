from docx import Document
from lxml import etree
import zipfile
from xml.sax.saxutils import escape, quoteattr
import re
ooXMLns = {'w':'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
#Function to extract all the comments of document(Same as accepted answer)
#Returns a dictionary with comment id as key and comment string as value

DEFAULT_ANNOTATION = {
    'application' : 'Working Copy Python Extractor',
}

class Comment():
    id: str
    commenttext: str
    author: str
    date : str
    redactionType : str
    exemptions : list = []
    checklist : list = []
    note : str

    startPara : int = -1
    startChar : int = -1
    startWord : str = ""
    endPara : int = -1
    endChar : int = -1
    endWord : str = ""
    runtext : str = ""

    def __init__(self, id, commenttext, author, date, redactionType, exemptions, checklist, note):
        self.id = id
        self.commenttext = commenttext
        self.author = author
        self.data = date
        self.redactionType = redactionType
        self.exemptions = exemptions
        self.checklist = checklist
        self.note = note

    def to_json(self, default):
        #rtr = default.copy()   INFO: removing additional fields, only 'extents' and 'applicationData' in 
        rtr = {}
        rtr['extents'] = [{
            'startPara' : self.startPara,
            'startChar' : self.startChar,
            'startWord' : self.startWord,
            'endPara' : self.endPara,
            'endChar' : self.endChar,
            'endWord' : self.endWord
        }]
        rtr['author'] = self.author
        rtr['applicationData'] = {
            #'run_text': self.runtext.strip(), 
            'text' : self.commenttext, 
            'redactionType' : self.redactionType,
            'exemptions' : self.exemptions,
            'checklist' : self.checklist,
            'note' : self.note
            #'sensitivity' : True, 
            #'entity' : False,
        }
        #print(rtr)
        return rtr


#cleanComment tidying up raw comment text for parsin (remove double spaces)
def cleanComment(comment):
    keys = ["Redaction:", "FOI-Exemptions:", "Checklist:", "Other Reason:", "Note:"]
    for key in keys:
        comment = comment.replace(key, " " + key + " ")
    comment = comment.replace("  ", " ")
    comment = comment.replace("Other Reason:", "Other-Reason:")
    return comment

#method for getting comment content from raw comment string
def get_redaction_exemptions(comment):
    comment = cleanComment(comment)

    #pattern to match " field: " (e.g. " redactionType: ", " exemptions: ")
    pattern = re.compile(' [a-zA-Z0-9\-]*?: ')
    matches = pattern.finditer(comment)

    #convert returned iterable to subscriptable list store [key, startIndex, endIndex]
    listed = []
    for match in matches:
        #get 
        temp = [str(match.group(0))[1:-2], int(match.span()[0]), int(match.span()[1])]
        listed = listed + [temp]

    #iterate through list retrieving key and associated value entry in comment, storing in new dict {key: textValue}
    limit = len(listed) - 1
    storageDict = {}
    i = 0
    while i < limit:
        key = listed[i][0]
        start = listed[i][2]
        end = listed[i+1][1]
        text = comment[start:end]

        #method of inserting checklist items in a list under 1 key in dict, while keeping other single entries as strings.
        if key in storageDict.keys():
            if type(storageDict[key]) == type([]):
                storageDict[key] = storageDict[key] + [text]
            else:
                storageDict[key] = [storageDict[key]] + [text]
        else:
            storageDict[key] = text
        

        i = i + 1

    key = listed[i][0]
    start = listed[i][2]
    text = comment[start:]
    if key in storageDict.keys():
        if type(storageDict[key]) == type([]):
            storageDict[key] = storageDict[key] + [text]
        else:
            storageDict[key] = [storageDict[key]] + [text]
    else:
        storageDict[key] = text


    #cleaning to convert exemptions numbers String into a List item based on comma splitting.
    if 'FOI-Exemptions' in storageDict.keys():
        storageDict['FOI-Exemptions'] = storageDict['FOI-Exemptions'].split(',')


    #union statement to provide empty values for any keys not present in original comment string.
    allKeys = ['Redaction','FOI-Exemptions','Checklist','Other-Reason','Note']
    currentKeys = list(storageDict.keys())
    non_intersect = list(set(currentKeys) ^ set(allKeys))
    for each in non_intersect:
        storageDict[each] = ""
    if 'FOI-Exemptions' in non_intersect:
        storageDict['FOI-Exemptions'] = []
    if 'Checklist' in non_intersect:
        storageDict['Checklist'] = []

    return storageDict

def get_document_comments(docxFileName):
    comments_dict={}
    docxZip = zipfile.ZipFile(docxFileName)

    try:
        commentsXML = docxZip.read('word/comments.xml')
    except KeyError:
        return []
    


    et = etree.XML(commentsXML)
    comments = et.xpath('//w:comment',namespaces=ooXMLns)
    for c in comments:
        commenttext=c.xpath('string(.)',namespaces=ooXMLns)
        comment_id=c.xpath('@w:id',namespaces=ooXMLns)[0]
        annotationDict = get_redaction_exemptions(commenttext)
        author = c.xpath('@w:author',namespaces=ooXMLns)[0]
        date = c.xpath('@w:date',namespaces=ooXMLns)[0]
        comment = Comment(comment_id, commenttext, author, date, annotationDict['Redaction'], annotationDict['FOI-Exemptions'], annotationDict['Checklist'], annotationDict['Note'])
        comments_dict[comment_id]=comment
    return comments_dict

def get_comment_start_ends(document, comments_dict):
    count_start = 0
    count_end = 0
    for para_index, paragraph in enumerate(document.paragraphs):
        for startElem in paragraph._element.xpath("./w:commentRangeStart"):
            comment_id = startElem.xpath('@w:id',namespaces=ooXMLns)[0]
            comment = comments_dict[comment_id]
            comment.startPara = para_index
            count_start += 1

        for endElem in paragraph._element.xpath("./w:commentRangeEnd"):
            comment_id = endElem.xpath('@w:id',namespaces=ooXMLns)[0]
            comment = comments_dict[comment_id]
            comment.endPara = para_index
            count_end += 1
    return (count_start, count_end)
    

def get_all_text(para) -> str:
    return para.text

def get_text_after(para, commentId, searchtag='Start') -> str:
    desiredStart = "./w:commentRange%s[@w:id='%s']" % (searchtag, commentId)
    printQueryText = desiredStart + "/following-sibling::*//text()"
    #print(printQueryText)
    return ''.join(para._element.xpath(printQueryText))

def get_text_before(para, commentId, searchtag='End') -> str:
    desiredStart = "./w:commentRange%s[@w:id='%s']" % (searchtag, commentId)
    printQueryText = desiredStart + "/preceding-sibling::*//text()"
    #print(printQueryText)
    return ''.join(para._element.xpath(printQueryText))

PARA_SEP = ' '


def getCCDHead(docxFileName, document):
    docTitleName = docxFileName.split("/")[-1][0:-5]
    docAuthor = document.core_properties.author
    docCreated = document.core_properties.created
    docAuthor = document.core_properties.author
    docLanguage = document.core_properties.language
    docLastModBy = document.core_properties.last_modified_by
    docLastModDate = document.core_properties.modified
    docRevision = document.core_properties.revision
    docSubject = document.core_properties.subject
    docTitle = document.core_properties.title

    heading = ''
    heading += '<ccd:doc xmlns:ccd="http://www.cicero.team/schema/ccd-20191205" srcfmt="word2003" docid=' + quoteattr(str(docTitleName)) +'>\n'
    heading += '<ccd:coreproperties>\n'
    heading += '<ccd:property name="created" value=' + quoteattr(str(docCreated))  +  '/>\n'
    heading += '<ccd:property name="creator" value=' + quoteattr(str(docAuthor)) + '/>\n'
    heading += '<ccd:property name="description" value=""/>'
    heading += '<ccd:property name="language" value=' + quoteattr(str(docLanguage)) + '/>\n'
    heading += '<ccd:property name="lastModifiedBy" value=' + quoteattr(str(docLastModBy)) + '/>\n'
    heading += '<ccd:property name="modified" value=' + quoteattr(str(docLastModDate)) + '/>\n'
    heading += '<ccd:property name="revision" value=' + quoteattr(str(docRevision)) + '/>\n'
    heading += '<ccd:property name="subject" value=' + quoteattr(str(docTitle)) + '/>\n'
    heading += '<ccd:property name="title" value=' + quoteattr(str(docTitle)) + '/>\n'
    heading += '</ccd:coreproperties>\n'
    return heading

def getCCDProperties(document):
    properties =  '<ccd:properties>\n'
    properties += '</ccd:properties>\n'
    return properties

def getCCDTail():

    tail = ''
    tail += '''<ccd:footer><ccd:para><ccd:field param='DOCPROPERTY "FolderReference"'>   </ccd:field></ccd:para>\n'''
    tail += '''<ccd:para><ccd:field param='DOCPROPERTY "Classification"'>   </ccd:field></ccd:para>\n'''
    tail += '''<ccd:para><ccd:field param='DOCPROPERTY "Caveat"'>   </ccd:field></ccd:para>\n'''
    tail += '''<ccd:para><ccd:field param='DOCPROPERTY "Descriptor"'>   </ccd:field></ccd:para>\n'''
    tail += '''<ccd:para><ccd:field param='DOCPROPERTY "Created"'>   </ccd:field></ccd:para>\n'''
    tail += '''<ccd:para>   </ccd:para>\n'''
    tail += '''</ccd:footer>\n'''
    tail +='''<ccd:header><ccd:para><ccd:field param='DOCPROPERTY "Classification"'>   </ccd:field></ccd:para>\n'''
    tail +='''<ccd:para><ccd:field param='DOCPROPERTY "Caveat"'>   </ccd:field></ccd:para>\n'''
    tail += '''<ccd:para><ccd:field param='DOCPROPERTY "Descriptor"'>   </ccd:field></ccd:para>\n'''
    tail += '''</ccd:header>\n'''
    tail += '''</ccd:section>\n'''
    tail += '''<ccd:metaproperties>\n'''
    tail += '''<ccd:property name="LCH" value="xx"/>\n'''
    tail += '''<ccd:property name="SCH" value="xx"/>\n'''
    tail += '''<ccd:property name="SAH" value="xx"/>\n'''
    tail += '''<ccd:property name="total-wordcount" value="xx"/>\n'''
    tail += '''<ccd:property name="content-wordcount" value="xx"/>\n'''
    tail += '''</ccd:metaproperties></ccd:doc>\n'''
    return tail


def getCCD(docxFileName : str):
    document = Document(docxFileName)
    heading = getCCDHead(docxFileName, document)
    properties = getCCDProperties(document)
    tail = getCCDTail()

    rtr = ""
    for para_index, paragraph in enumerate(document.paragraphs):
        if "DOCUMENT ATTRIBUTES SECTION: START" in paragraph.text:
            rtr += ('<ccd:section><ccd:body><ccd:para index="%d">%s</ccd:para>\n' % (para_index, escape(paragraph.text)))
        elif "DOCUMENT ATTRIBUTES SECTION: END" in paragraph.text:
            rtr += ("<ccd:para index='%d'>%s</ccd:para>\n" % (para_index, escape(paragraph.text)))
            #rtr += '<ccd:division name="Main Content"><ccd:para></ccd:para>\n'
        else:
            rtr += ("<ccd:para index='%d'>%s</ccd:para>\n" % (para_index, escape(paragraph.text)))

    rtr += '</ccd:body>\n'
    rtr = heading + properties + rtr + tail
    return rtr

def getAnnotations(docxFileName : str):
    document = Document(docxFileName)

    #TODO complete more of the JSON schema
    doc_default_annotation = DEFAULT_ANNOTATION.copy()
    doc_default_annotation['folioId'] = docxFileName

    comments_dict = get_document_comments(docxFileName)

    if comments_dict == []:
        return []

    (num_start, num_end) = get_comment_start_ends(document, comments_dict)

    assert num_start == len(comments_dict)
    assert num_end == len(comments_dict)
    
    for para_index, paragraph in enumerate(document.paragraphs):
        # print("<ccd:para index='%d'>%s</ccd:para>" % (para_index, paragraph.text))
        for c in comments_dict.values():
            paraCaptured = False

            #inside a comment
            if para_index > c.startPara and para_index < c.endPara:
                c.runtext += get_all_text(paragraph) + PARA_SEP

            # start or end para are handled specially
            elif para_index == c.startPara:
                before_comment_start_text = get_text_before(paragraph, c.id, searchtag='Start')
                c.startChar = len(before_comment_start_text)

                after_comment_start_text = get_text_after(paragraph, c.id)
                c.runtext += after_comment_start_text + PARA_SEP
                paraCaptured = True
                
            if para_index == c.endPara:
                text_before_end = get_text_before(paragraph, c.id, searchtag='End')
                if paraCaptured:
                    # this was also a start paragraph. we need to subtract any text captured AFTER the end tag
                    assert para_index == c.startPara
                    text_after_end = get_text_after(paragraph, c.id, searchtag='End')
                    if len(text_after_end):
                        c.runtext = c.runtext[:-len(text_after_end)]
                else:
                    c.runtext += text_before_end + PARA_SEP
                c.endChar = len(text_before_end)
            
            if len(c.runtext.split()) > 0:
                c.startWord = c.runtext.split()[0]
                c.endWord = c.runtext.split()[len(c.runtext.split()) - 1]
        

    return [c.to_json(doc_default_annotation) for c in comments_dict.values()]